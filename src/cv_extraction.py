import re
import fitz
from docx import Document
from dataclasses import dataclass
from datetime import datetime
from typing import List, Dict


SECTION_PATTERNS = {
    "skills": [
        r"(?i)^[\s]*(?:technical\s+)?skills[\s]*:?[\s]*$",
        r"(?i)^[\s]*core\s+(?:competencies|skills)[\s]*:?[\s]*$",
        r"(?i)^[\s]*(?:key\s+)?technologies[\s]*:?[\s]*$",
        r"(?i)^[\s]*technical\s+(?:proficiency|expertise)[\s]*:?[\s]*$",
        r"(?i)^[\s]*areas\s+of\s+expertise[\s]*:?[\s]*$",
        r"(?i)^[\s]*tools?\s*(?:&|and)?\s*technologies[\s]*:?[\s]*$",
        r"(?i)^[\s]*programming\s+(?:languages|skills)[\s]*:?[\s]*$",
    ],
    "summary": [
        r"(?i)^[\s]*(?:professional\s+)?summary[\s]*:?[\s]*$",
        r"(?i)^[\s]*(?:career\s+)?objective[\s]*:?[\s]*$",
        r"(?i)^[\s]*about\s*(?:me)?[\s]*:?[\s]*$",
        r"(?i)^[\s]*profile[\s]*:?[\s]*$",
        r"(?i)^[\s]*(?:executive\s+)?overview[\s]*:?[\s]*$",
        r"(?i)^[\s]*personal\s+statement[\s]*:?[\s]*$",
        r"(?i)^[\s]*introduction[\s]*:?[\s]*$",
    ],
    "experience": [
        r"(?i)^[\s]*(?:work\s+)?experience[\s]*:?[\s]*$",
        r"(?i)^[\s]*(?:professional\s+)?experience[\s]*:?[\s]*$",
        r"(?i)^[\s]*employment\s+(?:history|record)[\s]*:?[\s]*$",
        r"(?i)^[\s]*work\s+history[\s]*:?[\s]*$",
        r"(?i)^[\s]*career\s+history[\s]*:?[\s]*$",
    ],
    "education": [
        r"(?i)^[\s]*education[\s]*:?[\s]*$",
        r"(?i)^[\s]*academic\s+(?:background|qualifications)[\s]*:?[\s]*$",
        r"(?i)^[\s]*qualifications[\s]*:?[\s]*$",
    ],
    "projects": [
        r"(?i)^[\s]*projects?[\s]*:?[\s]*$",
        r"(?i)^[\s]*(?:key\s+)?projects?[\s]*:?[\s]*$",
        r"(?i)^[\s]*personal\s+projects?[\s]*:?[\s]*$",
    ],
    "certifications": [
        r"(?i)^[\s]*certifications?[\s]*:?[\s]*$",
        r"(?i)^[\s]*licenses?\s*(?:&|and)?\s*certifications?[\s]*:?[\s]*$",
        r"(?i)^[\s]*professional\s+certifications?[\s]*:?[\s]*$",
    ],
    "volunteering": [
        r"(?i)^[\s]*volunteer(?:ing)?(?:\s+(?:experience|work))?[\s]*:?[\s]*$",
        r"(?i)^[\s]*community\s+(?:service|involvement)[\s]*:?[\s]*$",
        r"(?i)^[\s]*(?:volunteer|civic)\s+activities[\s]*:?[\s]*$",
    ],
}


@dataclass
class CVData:
    raw_text: str
    skills: List[str]
    experience_years: float


def extract_text(file_path: str) -> str:
    def from_pdf():
        try:
            doc = fitz.open(file_path)
            text = ""

            for page in doc:
                text += page.get_text("text") + "\n"
            doc.close()
            
            return text.strip()
        except FileNotFoundError:
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {e}")

    def from_docx():
        try:
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        except FileNotFoundError:
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {e}")

    file_path_lower = file_path.lower()
    
    if file_path_lower.endswith('.pdf'):
        return from_pdf()
    elif file_path_lower.endswith('.docx'):
        return from_docx()
    else:
        raise ValueError("Unsupported format. Only .pdf and .docx are supported.")


def extract_sections(text: str) -> Dict[str, str]:
    if not text:
        return {}
    
    try:
        lines = text.split('\n')
        boundaries = []
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                continue
                
            for section_type, patterns in SECTION_PATTERNS.items():
                for pattern in patterns:
                    if re.match(pattern, line_stripped):
                        char_pos = sum(len(l) + 1 for l in lines[:i])
                        boundaries.append((section_type, i, char_pos))
                        break
        
        if not boundaries:
            return {}
        
        sections_content = {}
        
        sorted_boundaries = sorted(boundaries, key=lambda x: x[1])
        
        for i, (section_type, line_idx, _) in enumerate(sorted_boundaries):
            content_start = line_idx + 1
            
            if i + 1 < len(sorted_boundaries):
                content_end = sorted_boundaries[i + 1][1]
            else:
                content_end = len(lines)
            
            content_lines = lines[content_start:content_end]
            content = '\n'.join(line for line in content_lines if line.strip())
            
            if section_type in sections_content:
                sections_content[section_type] += '\n' + content.strip()
            else:
                sections_content[section_type] = content.strip()
        
        return sections_content
    except AttributeError:
        raise TypeError(f"Expected string input, got {type(text).__name__}")
    except Exception as e:
        raise ValueError(f"Failed to extract sections from text: {e}")


def parse_skills_from_section(skills_text: str) -> List[str]:
    if not skills_text:
        return []
    
    skills = []
    
    delimiters = r'[,;|•·▪▸►➤✓✔★●○◦\n\t]|\s{2,}|(?<=[a-z])(?=[A-Z])'
    
    category_pattern = r'([A-Za-z\s&/]+):\s*(.+?)(?=(?:[A-Za-z\s&/]+:|$))'
    category_matches = re.findall(category_pattern, skills_text, re.DOTALL)
    
    if category_matches:
        for category, skill_list in category_matches:
            parts = re.split(r'[,;|•·\n]+', skill_list)
            for part in parts:
                cleaned = part.strip()
                if cleaned and len(cleaned) > 1 and len(cleaned) < 50:
                    skills.append(cleaned)
    else:
        parts = re.split(delimiters, skills_text)
        for part in parts:
            cleaned = part.strip().strip('-').strip('•').strip()
            if cleaned and len(cleaned) > 1 and len(cleaned) < 50:
                if cleaned.count(' ') < 5:
                    skills.append(cleaned)
    
    seen = set()
    unique_skills = []
    for skill in skills:
        skill_lower = skill.lower()
        if skill_lower not in seen:
            seen.add(skill_lower)
            unique_skills.append(skill)
    
    return unique_skills


def extract_years_of_experience(summary_text: str, experience_text: str) -> float:
    def from_summary():
        if not summary_text:
            return None
        
        summary_lower = summary_text.lower()
        
        patterns = [
            r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s+(?:experience|exp)',
            r'(?:experience|exp)(?:\s+of)?\s+(\d+)\+?\s*(?:years?|yrs?)',
            r'worked\s+(?:for\s+)?(\d+)\+?\s*(?:years?|yrs?)',
            r'(\d+)\+?\s*(?:years?|yrs?)\s+(?:as\s+)?(?:a\s+)?[a-zA-Z\s]+(?:developer|engineer|analyst|scientist|manager)',
            r'over\s+(\d+)\s*(?:years?|yrs?)',
        ]
        
        total_years = 0
        for pattern in patterns:
            matches = re.findall(pattern, summary_lower, re.IGNORECASE)
            for match in matches:
                years = int(match) if isinstance(match, str) else int(match[0])
                total_years = max(total_years, years)
        
        return total_years if total_years > 0 else None

    def from_dates():
        if not experience_text:
            return None
        
        text_lower = experience_text.lower()
        
        date_pattern = r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)?\.?\s*(\d{4})\s*[-–—to]+\s*((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)?\.?\s*(\d{4})|present|current|now)'
        
        date_matches = re.findall(date_pattern, text_lower, re.IGNORECASE)
        
        if not date_matches:
            return None
        
        current_year = datetime.now().year
        job_periods = []
        
        for match in date_matches:
            start_year = int(match[0])
            end_part = match[1].strip()
            
            is_current = any(word in end_part.lower() for word in ['present', 'current', 'now'])
            
            if is_current:
                end_year = current_year
            else:
                year_match = re.search(r'(\d{4})', end_part)
                if year_match:
                    end_year = int(year_match.group(1))
                else:
                    continue
            
            if 1950 < start_year <= current_year and 1950 < end_year <= current_year and start_year <= end_year:
                job_periods.append((start_year, end_year))
        
        if not job_periods:
            return None
        
        job_periods.sort(key=lambda x: x[0])
        merged_periods = []
        
        for start, end in job_periods:
            if merged_periods and start <= merged_periods[-1][1]:
                merged_periods[-1] = (merged_periods[-1][0], max(merged_periods[-1][1], end))
            else:
                merged_periods.append((start, end))
        
        total_years = sum(end - start for start, end in merged_periods)
        
        return total_years if total_years > 0 else None

    summary_years = from_summary()
    if summary_years is not None:
        return summary_years
    
    exp_years = from_dates()
    return exp_years if exp_years is not None else 0


def extract_cv_data(file_path: str) -> CVData:
    if not file_path:
        raise ValueError("File path is required")
    
    raw_text = extract_text(file_path)
    
    sections_content = extract_sections(raw_text)
    
    skills_raw = sections_content.get("skills", "")
    skills = parse_skills_from_section(skills_raw)
    
    summary_section = sections_content.get("summary", "")
    experience_section = sections_content.get("experience", "")
    
    total_exp = extract_years_of_experience(summary_section, experience_section)
    
    return CVData(
        raw_text = raw_text,
        skills = skills,
        experience_years = total_exp
    )